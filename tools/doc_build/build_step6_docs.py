from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for name, width in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [("Heading 1", 15, RGBColor(46, 116, 181)), ("Heading 2", 12.5, RGBColor(31, 77, 120))]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(4)


def title(doc: Document, heading: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(heading)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def callout(doc: Document, label: str, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    margins(t)
    c = t.cell(0, 0)
    shade(c, "F4F6F9")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, rows: list[tuple[str, str]], header=("Part", "Purpose")) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    margins(t)
    h = t.rows[0].cells
    h[0].text = header[0]
    h[1].text = header[1]
    for c in h:
        shade(c, "E8EEF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        c = t.add_row().cells
        c[0].text = left
        c[1].text = right
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def technical() -> None:
    doc = Document()
    style(doc)
    title(doc, "Step 6 Technical Notes: One-Command Webcam Scan", "Embodied Visual Memory project - composing capture, detection, tracking, and memory preparation")
    callout(doc, "Objective", "Step 6 adds an orchestration command that runs the full webcam-to-memory pipeline in one command.")
    doc.add_heading("Pipeline", level=1)
    code(doc, "scan-webcam -> capture_webcam -> run_detection_for_dir -> run_tracking_for_dir -> queryable run folder")
    doc.add_heading("New Interfaces", level=1)
    table(doc, [
        ("scan-webcam", "Captures webcam frames, detects objects, draws annotations, tracks identities, and prints the next query command."),
        ("run_detection_for_dir", "Reusable helper used by both detect and scan-webcam."),
        ("run_tracking_for_dir", "Reusable helper used by both track and scan-webcam."),
    ])
    doc.add_heading("Math", level=1)
    doc.add_paragraph("Step 6 does not introduce new model math. It composes the previous math in a deterministic order:")
    bullets(doc, [
        "timestamp math from Step 1.",
        "bounding-box and confidence math from Step 2.",
        "IoU and center-distance tracking math from Step 3.",
        "memory retrieval readiness from Step 4.",
    ])
    doc.add_heading("What Was Built", level=1)
    bullets(doc, [
        "Added reusable detection and tracking helper functions in the CLI.",
        "Added scan-webcam command with capture, detector, annotation, and tracker options.",
        "Updated README with the one-command workflow.",
        "Verified an end-to-end run named smoke_pipeline.",
    ])
    doc.add_heading("Validation Performed", level=1)
    numbered(doc, [
        "Confirmed scan-webcam help output.",
        "Ran a 3-second scan named smoke_pipeline.",
        "Confirmed frames, observations, detections, annotations, tracks, and track summary were created.",
        "Confirmed list-memory works on the generated run.",
    ])
    doc.save(DOCS / "step_6_one_command_webcam_scan_technicalities_and_math.docx")


def simple() -> None:
    doc = Document()
    style(doc)
    title(doc, "Step 6 Simple Explanation: One-Command Webcam Scan", "Embodied Visual Memory project - making the project easier to use")
    callout(doc, "In plain English", "Step 6 turns the many separate commands into one command that records, understands, and prepares memory from the webcam.")
    doc.add_heading("What Changed From Step 5", level=1)
    doc.add_paragraph("Before this, you had to run capture, detect, track, and query commands separately. Step 6 gives you one command for the common webcam scan workflow.")
    doc.add_heading("What I Created In This Step", level=1)
    bullets(doc, [
        "A scan-webcam command.",
        "Shared internal helper functions so detect and track are not duplicated.",
        "A pipeline summary after the scan finishes.",
        "README instructions for the new easier workflow.",
    ])
    doc.add_heading("How You Use It", level=1)
    code(doc, "python -m evm.cli scan-webcam --run-name desk_scan --seconds 10")
    code(doc, "python -m evm.cli list-memory data\\runs\\desk_scan")
    code(doc, "python -m evm.cli query-memory data\\runs\\desk_scan bottle")
    doc.add_heading("Thought Process Behind The Decision", level=1)
    table(doc, [
        ("Reduce friction", "The project is easier to test when one command does the full scan."),
        ("Keep old commands", "Separate commands are still useful for debugging each step."),
        ("Reuse existing logic", "The one-command scan calls the same code as the manual workflow."),
        ("Print next steps", "After a scan, the terminal tells you what command to run next."),
    ])
    doc.add_heading("Important Limitation", level=1)
    doc.add_paragraph("This still uses the same object detector and tracker. The command is easier to use, but it does not make the AI more accurate by itself.")
    doc.save(DOCS / "step_6_one_command_webcam_scan_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    technical()
    simple()


if __name__ == "__main__":
    main()
