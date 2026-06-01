from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for name, width in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 58, 95)),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(4)


def title(doc: Document, heading: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(heading)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def callout(doc: Document, label: str, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    margins(t)
    c = t.cell(0, 0)
    shade(c, "F4F6F9")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, rows: list[tuple[str, str]], header=("Part", "Purpose")) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    margins(t)
    h = t.rows[0].cells
    h[0].text = header[0]
    h[1].text = header[1]
    for c in h:
        shade(c, "E8EEF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        c = t.add_row().cells
        c[0].text = left
        c[1].text = right
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def build_technical() -> None:
    doc = Document()
    style(doc)
    title(
        doc,
        "Step 3 Technical Notes: 2D Object Identity Tracking",
        "Embodied Visual Memory project - grouping detections into likely same-object tracks",
    )
    callout(
        doc,
        "Objective",
        "Step 3 turns repeated frame-by-frame detections into track IDs such as bottle_001 or person_001, creating the first version of object identity over time.",
    )

    doc.add_heading("Pipeline", level=1)
    code(doc, "detections.jsonl -> track association -> tracks.jsonl -> track_summary.json")

    doc.add_heading("New Interfaces", level=1)
    table(
        doc,
        [
            ("TrackRecord", "A detection record plus a stable track_id assigned by the tracker."),
            ("TrackSummary", "First seen, last seen, detection count, confidence, and last location for each tracked object."),
            ("track command", "Reads detections.jsonl and creates tracks.jsonl plus track_summary.json."),
            ("summarize-tracks command", "Prints each tracked object instance in a readable form."),
        ],
    )

    doc.add_heading("Tracking Math", level=1)
    doc.add_paragraph("The tracker compares each new detection with recent active tracks of the same label.")
    bullets(
        doc,
        [
            "Intersection over Union: IoU = overlap_area / union_area.",
            "Center distance: distance = sqrt((x2 - x1)^2 + (y2 - y1)^2).",
            "A detection can attach to a track if IoU is high enough or center distance is small enough.",
            "The default thresholds are IoU >= 0.25, center distance <= 180 pixels, and frame gap <= 10.",
        ],
    )
    doc.add_paragraph("The association score combines overlap and movement:")
    code(doc, "score = IoU + 0.25 * max(0, 1 - center_distance / max_center_distance)")

    doc.add_heading("What Was Built", level=1)
    bullets(
        doc,
        [
            "Added evm.tracking with bbox_iou, center_distance, track_detections, and summarize_tracks.",
            "Added CLI commands: track and summarize-tracks.",
            "Updated README with Step 3 usage.",
            "Verified Step 3 on the existing smoke_webcam run.",
        ],
    )

    doc.add_heading("Validation Performed", level=1)
    numbered(
        doc,
        [
            "Confirmed the new CLI commands are available.",
            "Ran tracking on data/runs/smoke_webcam after Step 2 detection.",
            "Generated tracks.jsonl and track_summary.json.",
            "Confirmed summarize-tracks prints object tracks with first/last timestamps.",
        ],
    )

    doc.add_heading("Known Limitation", level=1)
    doc.add_paragraph(
        "This is lightweight 2D tracking. Duplicate detector boxes can create duplicate tracks, and similar objects may still be confused. Later improvements can add stricter duplicate filtering, visual embeddings, or a proven tracker such as ByteTrack."
    )
    doc.save(DOCS / "step_3_object_identity_tracking_technicalities_and_math.docx")


def build_simple() -> None:
    doc = Document()
    style(doc)
    title(
        doc,
        "Step 3 Simple Explanation: Object Identity Tracking",
        "Embodied Visual Memory project - helping the bot know that this object is probably the same object over time",
    )
    callout(
        doc,
        "In plain English",
        "Step 3 helps the bot stop treating every detection as a brand-new object. It starts grouping repeated sightings into the same object track.",
    )

    doc.add_heading("What Changed From Step 2", level=1)
    doc.add_paragraph(
        "Step 2 could say that it saw a bottle in many frames. Step 3 tries to say those bottle sightings are probably the same bottle, so it creates an ID like bottle_001."
    )

    doc.add_heading("What I Created In This Step", level=1)
    bullets(
        doc,
        [
            "A tracker that reads the detected objects from Step 2.",
            "A track ID system like person_001, bottle_001, cup_001.",
            "A tracks file that records which detection belongs to which object track.",
            "A track summary file that says when each tracked object first and last appeared.",
        ],
    )

    doc.add_heading("Thought Process Behind The Decisions", level=1)
    table(
        doc,
        [
            ("Stay 2D for now", "Laptop cameras do not give real depth, so tracking image positions is the practical next move."),
            ("Use simple math first", "Overlap and movement are easy to understand and good enough for an early version."),
            ("Create track IDs", "Memory becomes more useful when the bot can refer to the same object over time."),
            ("Keep thresholds configurable", "Different rooms and camera angles may need different strictness."),
        ],
    )

    doc.add_heading("How You Check It", level=1)
    code(doc, "python -m evm.cli track data\\runs\\my_first_test")
    code(doc, "python -m evm.cli summarize-tracks data\\runs\\my_first_test")
    doc.add_paragraph("Then open the run folder and check:")
    bullets(doc, ["tracks.jsonl", "track_summary.json"])

    doc.add_heading("What The Math Means Simply", level=1)
    doc.add_paragraph(
        "If a box in the new frame overlaps strongly with a box from the previous frame, it is probably the same object. If it did not overlap much but the center moved only a little, it may still be the same object."
    )
    bullets(
        doc,
        [
            "Overlap tells us whether two boxes cover the same image region.",
            "Center distance tells us how far the object appears to move.",
            "Frame gap tells us how long the object was missing before we stop trusting the match.",
        ],
    )

    doc.add_heading("Important Limitation", level=1)
    doc.add_paragraph(
        "This is not perfect yet. If the detector creates duplicate boxes, the tracker may create duplicate tracks. That is normal at this stage. The goal is to build the first version of visual identity, then improve it."
    )

    doc.save(DOCS / "step_3_object_identity_tracking_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    build_technical()
    build_simple()


if __name__ == "__main__":
    main()
