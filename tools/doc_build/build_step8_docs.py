from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def setup(doc: Document) -> None:
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size in [("Heading 1", 15), ("Heading 2", 12.5)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)


def title(doc: Document, text: str, sub: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(sub)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def technical() -> None:
    doc = Document()
    setup(doc)
    title(doc, "Step 8 Technical Notes: HTML Visual Report", "Generating browsable visual evidence from tracked memory")
    doc.add_heading("Objective", level=1)
    doc.add_paragraph("Step 8 creates an HTML report from track summaries so results can be inspected visually without reading JSON.")
    doc.add_heading("Implementation", level=1)
    bullets(doc, [
        "Added evm.report with build_run_report.",
        "Added report CLI command.",
        "Each track becomes a card with image evidence and summary fields.",
        "Annotated frames are preferred when available; raw frames are used as fallback.",
    ])
    doc.add_heading("Math", level=1)
    doc.add_paragraph("Step 8 introduces no new CV math. It visualizes values already computed by detection and tracking: confidence, last frame, last center, and detection count.")
    doc.add_heading("Validation", level=1)
    numbered(doc, [
        "Generated report.html for smoke_pipeline.",
        "Confirmed the report file was created and references local evidence images.",
    ])
    doc.save(DOCS / "step_8_html_visual_report_technicalities_and_math.docx")


def simple() -> None:
    doc = Document()
    setup(doc)
    title(doc, "Step 8 Simple Explanation: HTML Visual Report", "A simple visual page for checking what the bot remembers")
    doc.add_heading("In Plain English", level=1)
    doc.add_paragraph("Step 8 creates a normal HTML page showing each remembered object with an image and useful details.")
    doc.add_heading("How You Use It", level=1)
    code(doc, "python -m evm.cli report data\\runs\\phone_room_scan")
    doc.add_paragraph("Then open report.html inside that run folder.")
    doc.add_heading("Why This Matters", level=1)
    bullets(doc, [
        "You can see what the bot remembers without opening JSON files.",
        "The evidence image makes wrong detections easier to spot.",
        "This is a first step toward a real app interface later.",
    ])
    doc.save(DOCS / "step_8_html_visual_report_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    technical()
    simple()


if __name__ == "__main__":
    main()
