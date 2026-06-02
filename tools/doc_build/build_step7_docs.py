from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def setup(doc: Document) -> None:
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size in [("Heading 1", 15), ("Heading 2", 12.5)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)


def title(doc: Document, text: str, sub: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(sub)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    tc_pr.append(shd)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def technical() -> None:
    doc = Document()
    setup(doc)
    title(doc, "Step 7 Technical Notes: Recorded Video Scan", "Processing phone-style video files through the visual memory pipeline")
    callout(doc, "Objective", "Step 7 lets prerecorded videos become observation runs, enabling phone-camera walkthroughs without live mobile streaming yet.")
    doc.add_heading("Pipeline", level=1)
    code(doc, "video file -> sampled frames -> observations.jsonl -> detection -> tracking -> queryable memory")
    doc.add_heading("Implementation", level=1)
    bullets(doc, [
        "Added VideoFileSource using OpenCV VideoCapture.",
        "Added frame sampling with --video-frame-stride.",
        "Added optional resizing with --resize-width.",
        "Added ingest-video for video-to-observations only.",
        "Added scan-video for full video-to-memory processing.",
    ])
    doc.add_heading("Math", level=1)
    bullets(doc, [
        "timestamp = source_frame_index / source_fps when FPS is available.",
        "resize scale = target_width / original_width.",
        "new_height = original_height * resize_scale.",
        "sampling keeps frames where source_frame_index modulo video_frame_stride equals zero.",
    ])
    doc.add_heading("Validation", level=1)
    numbered(doc, [
        "Created a small synthetic MP4 from existing captured frames.",
        "Ran scan-video on that MP4.",
        "Confirmed observations, detections, tracks, and memory were generated.",
    ])
    doc.save(DOCS / "step_7_recorded_video_scan_technicalities_and_math.docx")


def simple() -> None:
    doc = Document()
    setup(doc)
    title(doc, "Step 7 Simple Explanation: Recorded Video Scan", "Using phone-recorded videos as the bot's visual input")
    callout(doc, "In plain English", "Step 7 lets you record a video on your phone, put it in the project, and let the bot process it like a camera scan.")
    doc.add_heading("What Changed", level=1)
    doc.add_paragraph("Before this, the project mainly used the laptop webcam. Now it can process a saved video file, which is the first big move toward mobile-camera use.")
    doc.add_heading("How You Use It", level=1)
    code(doc, "python -m evm.cli scan-video data\\videos\\room_walkthrough.mp4 --run-name phone_room_scan")
    doc.add_heading("What It Creates", level=1)
    bullets(doc, ["frames", "observations.jsonl", "detections.jsonl", "annotated_frames", "tracks.jsonl", "track_summary.json"])
    doc.add_heading("Why This Matters", level=1)
    doc.add_paragraph("A phone camera can walk around a room more naturally than a laptop camera. This gives us better room scans without buying robot hardware.")
    doc.save(DOCS / "step_7_recorded_video_scan_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    technical()
    simple()


if __name__ == "__main__":
    main()
