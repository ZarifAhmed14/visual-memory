from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for name, width in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(31, 77, 120)),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(4)


def title(doc: Document, heading: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(heading)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def callout(doc: Document, label: str, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    margins(t)
    c = t.cell(0, 0)
    shade(c, "F4F6F9")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, rows: list[tuple[str, str]], header=("Part", "Purpose")) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    margins(t)
    h = t.rows[0].cells
    h[0].text = header[0]
    h[1].text = header[1]
    for c in h:
        shade(c, "E8EEF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        c = t.add_row().cells
        c[0].text = left
        c[1].text = right
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def technical() -> None:
    doc = Document()
    style(doc)
    title(doc, "Step 5 Technical Notes: Change Detection Between Runs", "Embodied Visual Memory project - comparing before and after visual memories")
    callout(doc, "Objective", "Step 5 compares two tracked runs and reports appeared, disappeared, persisted, and moved object labels.")

    doc.add_heading("Pipeline", level=1)
    code(doc, "before tracks + after tracks -> label set comparison + center shift -> change_report.json")

    doc.add_heading("New Interfaces", level=1)
    table(doc, [
        ("ChangeItem", "A single appeared, disappeared, persisted, or moved object entry with evidence paths."),
        ("ChangeReport", "Full before/after comparison result saved as JSON."),
        ("compare-runs", "CLI command that compares two run folders and writes change_report.json."),
        ("filter_tracks", "Removes weak tracks using detection count and confidence thresholds."),
    ])

    doc.add_heading("Math", level=1)
    bullets(doc, [
        "appeared = labels_after - labels_before.",
        "disappeared = labels_before - labels_after.",
        "persisted = labels_before intersect labels_after.",
        "center_shift = sqrt((after_x - before_x)^2 + (after_y - before_y)^2).",
        "moved if center_shift >= moved_distance_threshold, default 160 pixels.",
    ])

    doc.add_heading("What Was Built", level=1)
    bullets(doc, [
        "Added evm.change for label comparison, track filtering, and change report generation.",
        "Added compare-runs CLI command.",
        "Updated README with Step 5 usage.",
        "Verified by comparing smoke_webcam with itself.",
    ])

    doc.add_heading("Validation Performed", level=1)
    numbered(doc, [
        "Confirmed compare-runs help output.",
        "Compared data/runs/smoke_webcam against itself.",
        "Confirmed appeared and disappeared are empty.",
        "Confirmed person is persisted with center shift 0.0 pixels.",
        "Confirmed change_report.json is written.",
    ])

    doc.add_heading("Known Limitation", level=1)
    doc.add_paragraph("The comparison is label-level and image-space only. It does not yet know true 3D position or separate multiple same-label objects perfectly.")
    doc.save(DOCS / "step_5_change_detection_between_runs_technicalities_and_math.docx")


def simple() -> None:
    doc = Document()
    style(doc)
    title(doc, "Step 5 Simple Explanation: Change Detection Between Runs", "Embodied Visual Memory project - asking what changed between two camera scans")
    callout(doc, "In plain English", "Step 5 lets the bot compare a before scan and an after scan, then say what appeared, disappeared, stayed, or moved.")

    doc.add_heading("What Changed From Step 4", level=1)
    doc.add_paragraph("Step 4 answered questions about one run. Step 5 compares two runs, which makes the bot feel like it remembers changes over time.")

    doc.add_heading("What I Created In This Step", level=1)
    bullets(doc, [
        "A compare-runs command.",
        "A change report saved as change_report.json.",
        "Appeared, disappeared, still-present, and moved categories.",
        "Evidence paths pointing back to the before and after frames.",
    ])

    doc.add_heading("Thought Process Behind The Decisions", level=1)
    table(doc, [
        ("Compare labels first", "It is the clearest useful version of change detection."),
        ("Filter weak tracks", "Low-confidence detections can create fake changes."),
        ("Use center shift for movement", "Without depth, image movement is our practical movement estimate."),
        ("Save a report file", "Later UI or chatbot features can read the same structured output."),
    ])

    doc.add_heading("How You Check It", level=1)
    code(doc, "python -m evm.cli compare-runs data\\runs\\before_scan data\\runs\\after_scan")
    doc.add_paragraph("For a real test, record one scan, move or add an object, record a second scan, run detect and track on both, then compare them.")

    doc.add_heading("Important Limitation", level=1)
    doc.add_paragraph("This does not truly understand the room in 3D yet. It compares what the camera saw in two runs. Camera angle changes can affect the result.")
    doc.save(DOCS / "step_5_change_detection_between_runs_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    technical()
    simple()


if __name__ == "__main__":
    main()
