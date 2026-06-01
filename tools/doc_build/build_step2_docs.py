from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for margin_name, width in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 58, 95)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, rows: list[tuple[str, str]], header=("Part", "Purpose")) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    set_cell_margins(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for left, right in rows:
        cells = tbl.add_row().cells
        cells[0].text = left
        cells[1].text = right
    doc.add_paragraph()


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def build_technical() -> None:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Step 2 Technical Notes: Object Detection and Basic Visual Memory",
        "Embodied Visual Memory project - YOLO inference, detection records, annotated frames, and label memory",
    )

    add_callout(
        doc,
        "Objective",
        "Step 2 reads the saved observations from Step 1, runs object detection on the frame images, writes detection records, creates annotated evidence frames, and summarizes object memory by label.",
    )

    doc.add_heading("Pipeline", level=1)
    code_block(
        doc,
        "observations.jsonl -> load frame -> YOLO detector -> detections.jsonl\n"
        "                 -> annotated_frames/*.jpg\n"
        "                 -> memory_summary.json",
    )

    doc.add_heading("New Interfaces", level=1)
    table(
        doc,
        [
            ("DetectionRecord", "Frame-linked record containing label, confidence, bounding box, center point, area, image path, and model metadata."),
            ("YoloObjectDetector", "Thin wrapper around Ultralytics YOLO with a configurable model name and confidence threshold."),
            ("ObjectMemory", "Label-level summary with first seen, last seen, seen count, best confidence, average center, and average area."),
            ("detect command", "Runs inference over a saved run and writes detections, annotations, and memory summary."),
            ("summarize command", "Reads detections.jsonl and prints a human-readable object memory report."),
        ],
    )

    doc.add_heading("Detection Math", level=1)
    doc.add_paragraph("For every detected object, the model returns a class, confidence, and bounding box:")
    code_block(doc, "bbox = [x1, y1, x2, y2]")
    doc.add_paragraph("The implementation derives geometry used by memory and future tracking:")
    bullets(
        doc,
        [
            "width = max(0, x2 - x1)",
            "height = max(0, y2 - y1)",
            "center_x = x1 + width / 2",
            "center_y = y1 + height / 2",
            "area = width * height",
        ],
    )
    doc.add_paragraph(
        "Confidence thresholding keeps only detections above the configured minimum, defaulting to 0.35. YOLO internally applies non-maximum suppression to reduce duplicate overlapping boxes."
    )

    doc.add_heading("Memory Math", level=1)
    doc.add_paragraph("The first memory summary groups detections by label. For each label:")
    bullets(
        doc,
        [
            "first_seen is the minimum timestamp/frame for that label.",
            "last_seen is the maximum timestamp/frame for that label.",
            "seen_count is the number of detection records for that label.",
            "best_confidence is the maximum confidence score.",
            "average_center_xy is the mean of all detected object centers.",
            "average_area is the mean bounding-box area.",
        ],
    )
    doc.add_paragraph(
        "This is not full object identity tracking yet. Multiple bottles would currently be summarized under one bottle label. That is acceptable for Step 2 and becomes a target for the tracking step."
    )

    doc.add_heading("What Was Built", level=1)
    bullets(
        doc,
        [
            "Added ultralytics as the detection dependency.",
            "Added evm.detection for YOLO inference and annotated frame drawing.",
            "Added evm.memory for label-level visual memory summaries.",
            "Added CLI commands: detect and summarize.",
            "Updated README with Step 2 commands and output files.",
        ],
    )

    doc.add_heading("Validation Performed", level=1)
    numbered(
        doc,
        [
            "Installed the updated package in the project virtual environment.",
            "Verified detect and summarize command help output.",
            "Ran detection on data/runs/smoke_webcam with frame_stride=5.",
            "Confirmed detections.jsonl, annotated_frames, and memory_summary.json were created.",
            "Ran summarize and confirmed the smoke run reported person detections.",
        ],
    )

    doc.add_heading("Next Technical Step", level=1)
    doc.add_paragraph(
        "The next step should improve memory from label-level summaries to object identity over time, using center distance, box overlap, visual embeddings, or tracking algorithms."
    )

    doc.save(DOCS / "step_2_object_detection_and_basic_memory_technicalities_and_math.docx")


def build_simple() -> None:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Step 2 Simple Explanation: Object Detection and Basic Memory",
        "Embodied Visual Memory project - how the saved camera frames became object memories",
    )

    add_callout(
        doc,
        "In plain English",
        "Step 2 lets the bot look at the saved camera frames and name the things it sees, such as person, bottle, phone, laptop, or cup.",
    )

    doc.add_heading("What Changed From Step 1", level=1)
    doc.add_paragraph(
        "Step 1 only recorded camera frames. Step 2 analyzes those frames. The bot is no longer just saving images; it is starting to describe what appears inside them."
    )

    doc.add_heading("What I Created In This Step", level=1)
    bullets(
        doc,
        [
            "A detection command that reads a saved camera run.",
            "A pretrained YOLO object detector that finds common objects.",
            "A detections file that stores what was found in each frame.",
            "Annotated images that draw boxes around detected objects.",
            "A memory summary that says which object labels appeared and when they were last seen.",
        ],
    )

    doc.add_heading("Why I Chose This Approach", level=1)
    table(
        doc,
        [
            ("Use YOLO first", "It is practical, fast, well-known, and good enough for a laptop-camera v1."),
            ("Do not train yet", "Training would require a dataset. A pretrained model lets us build the system first."),
            ("Save detections to JSONL", "It keeps every detection easy to inspect and easy to feed into later memory steps."),
            ("Create annotated frames", "You can open the images and see if the model is right or wrong."),
            ("Summarize by label", "This gives us the first simple version of visual memory without overcomplicating tracking."),
        ],
    )

    doc.add_heading("How You Check It", level=1)
    doc.add_paragraph("After recording a run, use:")
    code_block(doc, "python -m evm.cli detect data\\runs\\my_first_test --frame-stride 5")
    doc.add_paragraph("Then show the memory summary:")
    code_block(doc, "python -m evm.cli summarize data\\runs\\my_first_test")
    doc.add_paragraph("Open the run folder and check:")
    bullets(
        doc,
        [
            "detections.jsonl for raw object records.",
            "memory_summary.json for the simple memory output.",
            "annotated_frames for images with boxes around detected objects.",
        ],
    )

    doc.add_heading("What The Math Means Simply", level=1)
    doc.add_paragraph(
        "The detector draws a rectangle around each object. From that rectangle, the system calculates the object's center and size. The confidence score is the model's strength of belief that the object label is correct."
    )
    bullets(
        doc,
        [
            "A bigger rectangle can mean the object is closer to the camera.",
            "The center tells us roughly where the object is in the image.",
            "The last timestamp tells us when the bot last saw that object label.",
            "The seen count tells us how often the object appeared in the run.",
        ],
    )

    doc.add_heading("What You Have To Do", level=1)
    numbered(
        doc,
        [
            "Record a new 10-20 second run with clear objects in view.",
            "Put common objects in front of the camera: bottle, cup, phone, laptop, keyboard, book, chair.",
            "Run the detect command.",
            "Open annotated_frames and see whether the boxes make sense.",
            "Run summarize and check which objects the bot remembers.",
        ],
    )

    doc.add_heading("Important Limitation", level=1)
    doc.add_paragraph(
        "Right now the memory is by object label, not by individual object identity. If there are two bottles, the system may summarize them together as bottle. Later we will add tracking so it can separate object instances."
    )

    doc.save(DOCS / "step_2_object_detection_and_basic_memory_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    build_technical()
    build_simple()


if __name__ == "__main__":
    main()
