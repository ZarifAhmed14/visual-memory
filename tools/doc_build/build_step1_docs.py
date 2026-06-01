from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for margin_name, width in {
        "top": "80",
        "bottom": "80",
        "start": "120",
        "end": "120",
    }.items():
        node = cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 58, 95)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 58, 95)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths=(2.0, 4.3)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_width(table, list(widths))
    hdr = table.rows[0].cells
    hdr[0].text = "Part"
    hdr[1].text = "Purpose"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    doc.add_paragraph()


def build_technical() -> None:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Step 1 Technical Notes: Visual Input Pipeline",
        "Embodied Visual Memory project - webcam capture, observation records, storage, and replay",
    )

    add_callout(
        doc,
        "Objective",
        "Step 1 creates the observation stream: a reliable way to turn webcam frames into timestamped, replayable data that later detection and memory modules can consume.",
    )

    doc.add_heading("System Contract", level=1)
    doc.add_paragraph(
        "The pipeline is intentionally model-free. It does not detect objects yet. Its job is to produce stable visual observations with frame IDs, timestamps, image paths, source metadata, and optional future fields for depth, camera pose, and intrinsics."
    )
    add_key_value_table(
        doc,
        [
            ("WebcamSource", "Opens the laptop camera through OpenCV, captures BGR frames, assigns frame IDs, and records elapsed timestamps."),
            ("FramePacket", "In-memory packet containing frame_id, timestamp, raw frame array, and metadata such as width, height, and camera index."),
            ("RunWriter", "Writes frames to disk and appends one JSON observation record per frame."),
            ("ObservationRecord", "Serializable metadata record used as the long-term interface between input, replay, and future CV modules."),
            ("Replay command", "Reads observations.jsonl, loads each image from disk, and displays the run in order for debugging."),
        ],
    )

    doc.add_heading("Data Model", level=1)
    doc.add_paragraph("Each saved observation has this logical shape:")
    code = doc.add_paragraph()
    code.style = doc.styles["Normal"]
    run = code.add_run(
        "frame_id: int\n"
        "timestamp: float\n"
        "rgb_path: str\n"
        "depth_path: str | None\n"
        "camera_pose: list[list[float]] | None\n"
        "camera_intrinsics: dict | None\n"
        "source: str\n"
        "metadata: dict"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("Timestamp Math", level=1)
    doc.add_paragraph(
        "For the webcam v1, timestamp is measured as elapsed monotonic time from the moment the camera source opens:"
    )
    p = doc.add_paragraph()
    r = p.add_run("timestamp_seconds = perf_counter_now - perf_counter_at_source_open")
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    doc.add_paragraph(
        "This avoids wall-clock jumps and gives a consistent timeline for later memory events. Frame rate is not assumed to be constant; each frame carries its own timestamp."
    )

    doc.add_heading("Storage Layout", level=1)
    doc.add_paragraph("A capture run is stored under data/runs/<run-name>:")
    add_bullets(
        doc,
        [
            "frames/000001.jpg, frames/000002.jpg, ... store the visual evidence.",
            "observations.jsonl stores one metadata object per line, which is easy to stream for long recordings.",
            "manifest.json stores run-level summary fields such as source, start time, frame count, and relative paths.",
        ],
    )

    doc.add_heading("Implementation Decisions", level=1)
    add_key_value_table(
        doc,
        [
            ("Python", "Fast to prototype, strong CV ecosystem, easy for later OpenCV and model integration."),
            ("OpenCV", "Standard practical library for webcam capture, image encoding, preview windows, and replay."),
            ("JSONL", "Append-friendly, human-readable, and scalable enough for frame-by-frame observation logs."),
            ("Relative frame paths", "Keeps runs movable inside the project folder without rewriting metadata."),
            ("Optional depth and pose fields", "Preserves the future robot/phone/AR interface without requiring hardware today."),
        ],
    )

    doc.add_heading("What Was Built", level=1)
    add_bullets(
        doc,
        [
            "Created project folder D:\\embodied-visual-memory, fully isolated on the D: drive.",
            "Added a src-based Python package named evm.",
            "Added capture-webcam command for live camera recording.",
            "Added replay command for reading saved runs.",
            "Created a local virtual environment and installed the package in editable mode.",
            "Verified webcam capture with a smoke run that saved 76 frames at 1280x720.",
        ],
    )

    doc.add_heading("Validation Performed", level=1)
    add_numbered(
        doc,
        [
            "Confirmed Python 3.11 and Git are available.",
            "Installed project dependencies in D:\\embodied-visual-memory\\.venv.",
            "Ran CLI help/import smoke checks.",
            "Captured a 3-second webcam run named smoke_webcam.",
            "Inspected manifest.json, observations.jsonl, and saved frame count.",
            "Ran replay against the saved smoke run.",
        ],
    )

    doc.add_heading("Next Technical Step", level=1)
    doc.add_paragraph(
        "Step 2 should consume ObservationRecord entries and run object detection over rgb_path frames, producing Detection records linked back to frame_id and timestamp."
    )

    doc.save(DOCS / "step_1_visual_input_pipeline_technicalities_and_math.docx")


def build_plain() -> None:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Step 1 Simple Explanation: Visual Input Pipeline",
        "Embodied Visual Memory project - what we made, why we made it, and how to use it",
    )

    add_callout(
        doc,
        "In plain English",
        "We built the bot's eyes and recorder. It can look through the laptop camera, save what it sees, mark each frame with a time, and replay the recording later.",
    )

    doc.add_heading("What We Are Trying To Build", level=1)
    doc.add_paragraph(
        "The bigger project is a visual-memory bot. Instead of only remembering text, it will remember visual evidence from a camera: objects, when they appeared, and eventually what changed."
    )
    doc.add_paragraph(
        "Step 1 does not understand objects yet. It simply creates the clean camera history that future steps will analyze."
    )

    doc.add_heading("What I Created In This Step", level=1)
    add_bullets(
        doc,
        [
            "A project folder on the D: drive only: D:\\embodied-visual-memory.",
            "A webcam capture command that records short camera sessions.",
            "A folder structure that saves every captured frame.",
            "A metadata file that records frame number, time, image path, camera size, and camera index.",
            "A replay command so we can inspect what the system captured.",
            "A private Python environment so this project does not depend on random system packages.",
        ],
    )

    doc.add_heading("Why This Step Matters", level=1)
    doc.add_paragraph(
        "If the camera input is messy, every future AI feature becomes painful. Object detection, memory, and change detection all need a clean stream of frames. This step makes the future work easier because every frame is saved in a predictable format."
    )

    doc.add_heading("Thought Process Behind The Decisions", level=1)
    add_key_value_table(
        doc,
        [
            ("Start with laptop webcam", "It costs nothing and lets us prove the idea before using a phone or robot camera."),
            ("Save frames to disk", "Future object detection needs actual image files to inspect, debug, and compare."),
            ("Use timestamps", "Memory needs time. Without timestamps, the bot cannot say when it last saw something."),
            ("Use JSONL metadata", "It is simple, readable, and works well when there are many frames."),
            ("Keep depth/pose optional", "We can add phone camera or AR data later without redesigning the whole system."),
        ],
    )

    doc.add_heading("How You Use It", level=1)
    doc.add_paragraph("Open PowerShell and run:")
    p = doc.add_paragraph()
    r = p.add_run(
        "cd D:\\embodied-visual-memory\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "python -m evm.cli capture-webcam --run-name my_first_test --seconds 10"
    )
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    doc.add_paragraph("Then replay the saved camera run:")
    p = doc.add_paragraph()
    r = p.add_run("python -m evm.cli replay data\\runs\\my_first_test")
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    doc.add_heading("What You Have To Do", level=1)
    add_numbered(
        doc,
        [
            "Place a few simple objects in front of the laptop camera, such as a phone, bottle, cup, notebook, or keys.",
            "Run a 10-second capture.",
            "Move one or two objects during the capture if you want more interesting test data.",
            "Replay the result and confirm the saved frames look clear.",
        ],
    )

    doc.add_heading("What Comes Next", level=1)
    doc.add_paragraph(
        "The next step is object detection. That is when the bot starts turning frames into labels like bottle, cup, phone, keyboard, and notebook. After that, we can store those labels as memory."
    )

    doc.save(DOCS / "step_1_visual_input_pipeline_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    build_technical()
    build_plain()


if __name__ == "__main__":
    main()
