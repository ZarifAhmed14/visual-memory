from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"
OUT = DOCS / "embodied_visual_memory_project_teaching_guide.docx"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for name, width in {"top": "100", "bottom": "100", "start": "130", "end": "130"}.items():
        node = cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 58, 95)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Embodied Visual Memory Project Teaching Guide")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("How the vision, motion, memory, commands, outputs, and math work")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project folder: D:\\embodied-visual-memory")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(85, 85, 85)
    doc.add_paragraph()


def callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    margins(table)
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    margins(t)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade(hdr[i], "E8EEF5")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    title(doc)

    callout(
        doc,
        "One-sentence explanation",
        "This project is a visual-memory bot: it takes webcam or phone video, detects objects, tracks them over time, remembers when they were last seen, compares scans, and creates a visual report.",
    )

    doc.add_heading("1. What This Project Can Achieve", level=1)
    bullets(
        doc,
        [
            "Convert laptop webcam or phone video into saved visual observations.",
            "Detect common objects such as cup, bottle, chair, book, phone, person, and bed using YOLO.",
            "Draw boxes around detected objects so humans can inspect what the AI believed it saw.",
            "Group repeated detections into rough object tracks such as cup_001 or chair_002.",
            "Answer simple memory questions like where it last saw a bottle.",
            "Compare a before scan and an after scan to report appeared, disappeared, persisted, and moved objects.",
            "Generate an HTML report with object cards and evidence images.",
        ],
    )

    doc.add_heading("2. Big Picture Pipeline", level=1)
    code(
        doc,
        "webcam or phone video\n"
        "-> frames + timestamps\n"
        "-> object detection\n"
        "-> annotated evidence frames\n"
        "-> object tracking\n"
        "-> visual memory\n"
        "-> query, compare, report",
    )
    doc.add_paragraph(
        "The important idea is that every step produces files that the next step can read. This makes the system easier to debug than a single black-box AI script."
    )

    doc.add_heading("3. How Vision Works In This Project", level=1)
    doc.add_paragraph(
        "Vision starts with pixels. The system captures video frames and gives each frame a timestamp. A YOLO object detector then looks at each frame and predicts object labels, confidence scores, and bounding boxes."
    )
    table(
        doc,
        ["Input", "Processing", "Expected outcome"],
        [
            ["Raw webcam frame", "Save image and timestamp", "frames/000001.jpg plus an observation record"],
            ["Saved image frame", "YOLO object detection", "labels like cup, chair, bottle with confidence scores"],
            ["Detected object", "Draw bounding box", "annotated_frames image showing visual evidence"],
            ["Many frames", "Repeat detection", "detections.jsonl with object records across time"],
        ],
        [1.8, 2.3, 3.0],
    )

    doc.add_heading("4. How Motion And Tracking Work", level=1)
    doc.add_paragraph(
        "The project does not have true 3D motion yet. It uses 2D image-space motion. If an object appears in nearby positions across frames, the tracker treats it as probably the same object."
    )
    bullets(
        doc,
        [
            "Same label: a cup can only match another cup track, not a chair track.",
            "Bounding-box overlap: if two boxes cover the same area, they are likely the same object.",
            "Center distance: if the object center moves only a small amount, it may still be the same object.",
            "Frame gap: if too many frames pass, the tracker becomes less willing to connect detections.",
        ],
    )
    callout(
        doc,
        "Important limitation",
        "This is rough 2D tracking, not perfect object permanence. If the camera moves quickly or YOLO creates duplicate boxes, the tracker can create duplicate tracks.",
    )

    doc.add_heading("5. Where The Math Is And What It Does", level=1)
    table(
        doc,
        ["Math location", "Formula or idea", "What it does"],
        [
            ["sources.py", "timestamp = source_frame_index / FPS", "Gives every video frame a time position so memory can say when an object was seen."],
            ["sources.py", "resize scale = target_width / original_width", "Shrinks large phone videos while preserving aspect ratio."],
            ["detection.py", "width = x2 - x1, height = y2 - y1", "Measures the size of the detected object box."],
            ["detection.py", "center = (x1 + width/2, y1 + height/2)", "Finds the object position inside the image."],
            ["detection.py", "area = width * height", "Estimates how much image space the object covers."],
            ["tracking.py", "IoU = overlap area / union area", "Measures whether two boxes probably refer to the same object."],
            ["tracking.py", "distance = sqrt((x2-x1)^2 + (y2-y1)^2)", "Measures how far the object center moved between frames."],
            ["tracking.py", "association score = IoU + distance bonus", "Chooses the best previous track for a new detection."],
            ["memory.py", "mean center, mean area, max confidence", "Summarizes repeated detections into useful memory statistics."],
            ["change.py", "after labels - before labels", "Finds appeared objects."],
            ["change.py", "before labels - after labels", "Finds disappeared objects."],
            ["change.py", "center shift distance", "Flags an object as moved if its image position changed enough."],
            ["query.py", "label similarity score", "Matches user words like phone to detector labels like cell phone."],
        ],
        [1.55, 2.25, 3.2],
    )

    doc.add_heading("6. What Each Code Part Does", level=1)
    table(
        doc,
        ["File", "What it controls", "Expected outcome"],
        [
            ["observation.py", "Defines the saved observation record.", "Each frame has frame_id, timestamp, image path, source, and metadata."],
            ["sources.py", "Reads webcam streams and video files.", "Turns camera/video input into frame packets."],
            ["storage.py", "Writes and reads run folders.", "Creates frames, observations.jsonl, and manifest.json."],
            ["detection.py", "Runs YOLO and draws object boxes.", "Creates detection records and annotated evidence images."],
            ["memory.py", "Summarizes detections by label.", "Creates memory_summary.json with first seen, last seen, and confidence."],
            ["tracking.py", "Groups detections into object tracks.", "Creates tracks.jsonl and track_summary.json."],
            ["query.py", "Searches tracked memory by object name.", "Answers where an object was last seen."],
            ["change.py", "Compares two runs.", "Creates change_report.json with appeared/disappeared/moved/persisted objects."],
            ["report.py", "Creates an HTML report.", "Creates report.html with visual memory cards."],
            ["cli.py", "Connects all features to terminal commands.", "Lets users run scan-video, scan-webcam, query-memory, compare-runs, and report."],
        ],
        [1.5, 2.35, 3.15],
    )

    doc.add_heading("7. Commands And Expected Outcomes", level=1)
    table(
        doc,
        ["Command", "What it does", "Expected outcome"],
        [
            ["scan-webcam --run-name desk_scan", "Captures webcam, detects objects, tracks them.", "A run folder named desk_scan with frames, detections, tracks, and summaries."],
            ["scan-video video.mp4 --run-name phone_scan", "Processes recorded phone video.", "A queryable memory run from the video."],
            ["list-memory data/runs/phone_scan", "Lists remembered object tracks.", "Text output like cup_001 last seen at 11.57s."],
            ["query-memory data/runs/phone_scan bottle", "Looks for a remembered object.", "Answer with track ID, last timestamp, frame number, and evidence path."],
            ["compare-runs before after", "Compares two scans.", "Appeared, disappeared, still present, and moved categories."],
            ["report data/runs/phone_scan", "Builds visual report.", "report.html inside the run folder."],
        ],
        [2.45, 2.1, 2.45],
    )

    doc.add_heading("8. How To Read The Report", level=1)
    bullets(
        doc,
        [
            "track_id is the memory ID, for example cup_001.",
            "label is what the detector thinks the object is.",
            "last seen tells when the object last appeared in the video.",
            "detections tells how many times the object was detected.",
            "best confidence tells how confident the model was at its strongest moment.",
            "the image is the evidence frame; check whether the box actually points to the object.",
        ],
    )
    callout(
        doc,
        "How to judge correctness",
        "Label correct plus box correct means good memory. Label wrong but box on a real object means the detector guessed the wrong category. A box around nothing useful is a false detection.",
    )

    doc.add_heading("9. What Your Group Needs To Learn", level=1)
    numbered(
        doc,
        [
            "Python basics: functions, classes, dataclasses, file paths, JSON, and command-line arguments.",
            "OpenCV basics: reading video, capturing webcam frames, saving images, drawing boxes.",
            "Computer vision concepts: frames, pixels, object detection, bounding boxes, confidence scores.",
            "YOLO basics: pretrained object detector, labels, confidence threshold, model limitations.",
            "Tracking basics: IoU, center distance, frame gaps, and why object identity is hard.",
            "Data design: why the project saves observations, detections, tracks, summaries, and reports separately.",
            "Evaluation mindset: inspect evidence images instead of blindly trusting model text.",
            "Portfolio storytelling: explain the problem, pipeline, demo, limitations, and future improvements.",
        ],
    )

    doc.add_heading("10. What This Project Does Not Do Yet", level=1)
    bullets(
        doc,
        [
            "It does not know true 3D position because the laptop/phone video does not provide reliable depth.",
            "It does not perform perfect long-term object identity when multiple similar objects exist.",
            "It does not understand full natural-language questions yet; object lookup is simple.",
            "It does not have a polished frontend/backend product interface yet.",
            "It does not control a real robot yet.",
        ],
    )

    doc.add_heading("11. Teaching Flow For Your Group", level=1)
    numbered(
        doc,
        [
            "Start with the one-sentence explanation: a bot that sees, remembers, and reports visual evidence.",
            "Show the pipeline diagram from raw video to report.",
            "Run or show scan-video on the mobile clip.",
            "Open report.html and explain each card.",
            "Explain the math using boxes, centers, overlap, and distance.",
            "Explain each source file by responsibility, not by reading code line by line.",
            "Discuss limitations honestly and present future upgrades: web UI, better tracking, RAG-style natural-language query, and eventually mobile live camera.",
        ],
    )

    doc.add_heading("12. Best Demo Script", level=1)
    numbered(
        doc,
        [
            "Show a 20-second phone video of a room or desk.",
            "Run scan-video and show the terminal summary: frames, detections, tracks.",
            "Run list-memory and show remembered objects.",
            "Run query-memory for one object.",
            "Open report.html and inspect object cards.",
            "Explain one correct detection and one wrong detection to show mature evaluation.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
