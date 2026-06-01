from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for name, width in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            cell_mar.append(node)
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")


def style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(31, 77, 120)),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(4)


def title(doc: Document, heading: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(heading)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def callout(doc: Document, label: str, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    margins(t)
    c = t.cell(0, 0)
    shade(c, "F4F6F9")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, rows: list[tuple[str, str]], header=("Part", "Purpose")) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    margins(t)
    h = t.rows[0].cells
    h[0].text = header[0]
    h[1].text = header[1]
    for c in h:
        shade(c, "E8EEF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        c = t.add_row().cells
        c[0].text = left
        c[1].text = right
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def technical() -> None:
    doc = Document()
    style(doc)
    title(doc, "Step 4 Technical Notes: Queryable Visual Memory", "Embodied Visual Memory project - label matching, last-seen retrieval, and grounded answers")
    callout(doc, "Objective", "Step 4 adds a user-facing query layer over tracked objects so the system can list remembered tracks and answer where a target object was last seen.")

    doc.add_heading("Pipeline", level=1)
    code(doc, "tracks.jsonl -> summarize_tracks -> label matching -> answer + supporting frame path")

    doc.add_heading("New Interfaces", level=1)
    table(doc, [
        ("MemoryQueryResult", "Structured result with query text, match flag, score, chosen track, supporting frame path, and answer text."),
        ("list-memory", "Prints every tracked memory with last-seen timestamp and evidence frame."),
        ("query-memory", "Finds the best object-label match and answers with the latest supporting visual evidence."),
        ("normalize_label", "Handles simple aliases such as phone -> cell phone."),
    ])

    doc.add_heading("Math and Matching", level=1)
    bullets(doc, [
        "Exact normalized label match scores 1.0.",
        "Substring label match scores 0.85.",
        "Fuzzy fallback uses SequenceMatcher ratio between query label and remembered label.",
        "Default minimum match score is 0.55.",
        "If multiple tracks match, the implementation chooses by score and latest last-seen timestamp.",
    ])

    doc.add_heading("What Was Built", level=1)
    bullets(doc, [
        "Added evm.query for query normalization, fuzzy label matching, and answer creation.",
        "Added CLI commands: list-memory and query-memory.",
        "Updated README with Step 4 usage.",
        "Verified positive and negative queries against the smoke_webcam run.",
    ])

    doc.add_heading("Validation Performed", level=1)
    numbered(doc, [
        "Confirmed list-memory and query-memory help output.",
        "Listed remembered tracks from data/runs/smoke_webcam.",
        "Queried person and received a last-seen answer with frame evidence.",
        "Queried bottle and received a clear not-found answer.",
    ])

    doc.add_heading("Known Limitation", level=1)
    doc.add_paragraph("This is not natural-language question answering yet. It is object-label retrieval. A later chatbot layer can parse richer questions and call this memory query code underneath.")
    doc.save(DOCS / "step_4_queryable_visual_memory_technicalities_and_math.docx")


def simple() -> None:
    doc = Document()
    style(doc)
    title(doc, "Step 4 Simple Explanation: Queryable Visual Memory", "Embodied Visual Memory project - asking the bot what it remembers")
    callout(doc, "In plain English", "Step 4 lets you ask the saved visual memory about an object and get an answer with the last frame where it was seen.")

    doc.add_heading("What Changed From Step 3", level=1)
    doc.add_paragraph("Step 3 created object tracks, but you still had to inspect files. Step 4 gives you simple commands to ask what the system remembers.")

    doc.add_heading("What I Created In This Step", level=1)
    bullets(doc, [
        "A command to list all remembered object tracks.",
        "A command to ask about one object.",
        "Simple matching so words like phone can match cell phone.",
        "Answers that include timestamp, frame number, and supporting image path.",
    ])

    doc.add_heading("Thought Process Behind The Decisions", level=1)
    table(doc, [
        ("Start with object names", "This is easier and more reliable than full chatbot questions at this stage."),
        ("Return evidence paths", "The answer should be grounded in an actual saved frame, not just text."),
        ("Use fuzzy matching", "Users may type slightly different words than the detector label."),
        ("Keep it CLI first", "The memory logic should work before we wrap it in a nicer app interface."),
    ])

    doc.add_heading("How You Check It", level=1)
    code(doc, "python -m evm.cli list-memory data\\runs\\my_first_test")
    code(doc, "python -m evm.cli query-memory data\\runs\\my_first_test bottle")

    doc.add_heading("What You Should See", level=1)
    bullets(doc, [
        "If the object was seen, the system tells you the track ID, last timestamp, frame number, and image path.",
        "If the object was not seen, the system says it could not find it and lists what it does remember.",
    ])

    doc.add_heading("Important Limitation", level=1)
    doc.add_paragraph("This is still not a full conversation system. It answers object lookup questions. Later we can make it understand questions like 'what changed?' or 'where was my phone before?'")
    doc.save(DOCS / "step_4_queryable_visual_memory_simple_explanation.docx")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    technical()
    simple()


if __name__ == "__main__":
    main()
